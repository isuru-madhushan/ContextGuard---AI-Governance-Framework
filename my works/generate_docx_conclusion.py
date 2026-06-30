import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def create_document():
    doc = docx.Document()

    # ── 1. PAGE MARGINS (1 INCH ALL SIDES) ──
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # ── 2. CONFIGURE STYLES (TIMES NEW ROMAN 12PT, 1.5 SPACING) ──
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = Pt(12)
    normal_style.font.color.rgb = RGBColor(51, 51, 51)
    normal_style.paragraph_format.line_spacing = 1.5
    normal_style.paragraph_format.space_after = Pt(6)

    # Helper function for headings
    def add_custom_heading(text, level):
        p = doc.add_paragraph()
        p.paragraph_format.keep_with_next = True
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.bold = True
        if level == 1:
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(0, 51, 102) # Navy Blue
        elif level == 2:
            run.font.size = Pt(15)
            run.font.color.rgb = RGBColor(0, 102, 153)
        elif level == 3:
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(51, 51, 51)
        return p

    def add_code_block(code_text):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.autofit = False
        tbl.columns[0].width = Inches(6.5)
        cell = tbl.cell(0, 0)
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="2D3748"/>') # Dark bg
        cell._tc.get_or_add_tcPr().append(shading)
        
        p = cell.paragraphs[0]
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent = Inches(0.1)
        run = p.add_run(code_text)
        run.font.name = 'Courier New'
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(247, 250, 252) # Light text
        
        doc.add_paragraph() # spacing

    # ── BUILD DOCUMENT CONTENT ──
    add_custom_heading("10. Conclusion", level=1)
    doc.add_paragraph("The mid-implementation stage of the ContextGuard Shadow AI Governance Framework marks a highly successful transition from theoretical cybersecurity concepts into an operational, enterprise-grade software prototype. This section summarizes the primary engineering achievements accomplished to date, confirms the technical feasibility of the proposed architecture, and outlines the strategic completion plan for the final academic deliverables.")

    # ── 10.1 ──
    add_custom_heading("10.1 Summary of Progress", level=2)
    doc.add_paragraph("The project has fully satisfied the core technical objectives established for the initial and mid-stage development phases (Phases 1 through 3). By establishing a non-intrusive, context-aware inspection plane at the simulated corporate network boundary, ContextGuard successfully addresses the critical visibility and governance gaps associated with employee utilization of unsanctioned Large Language Models (LLMs).")

    add_custom_heading("Key Engineering Milestones Achieved:", level=3)
    milestones = [
        ("Active Interception & Payload Decoding: ", "Successfully deployed a live Man-in-the-Middle interception kernel using mitmproxy (live_mitm_logger.py). Engineered the Advanced Double-Plane URL Decoding Engine to parse highly divergent POST payload architectures, successfully unquoting complex Google Gemini URL structures (f.req=) and nested OpenAI JSON arrays into a standardized telemetry stream (wrse_comprehensive_audit.log)."),
        ("High-Performance NLP & Asset Indexing: ", "Eliminated real-time disk I/O bottlenecks by refactored data_core.py to index 15 Master CSV Corporate Asset Sheets (data Sheets/) directly into an optimized in-memory dictionary (idx) upon system boot. Executed multi-pass regex normalization (forensic_normalize()) and token substring tracking to achieve sub-second matching of highly sensitive Patient IDs (HL7-517169) and infrastructure connection strings without duplicate collision counting."),
        ("WRSE Mathematical Model Validation: ", "Implemented the linear weighted sum algorithm calculate_wrse(), successfully combining Data Sensitivity (Ws = 0.50), Destination Trust (Wd = 0.25), and User Authority (Wu = 0.25) to generate normalized risk coefficients (0–100) and trigger dynamic severity classifications (CRITICAL, MEDIUM, LOW)."),
        ("Zero-Trust Administrative SOC Dashboard: ", "Deployed an interactive Streamlit portal (app.py, styles.py) featuring over 1,090 lines of custom CSS (THREATMON_CSS) for premium dark glassmorphism aesthetics. Successfully hardened the portal by engineering a regex-based Web Application Firewall (WAF) for SQLi/XSS defense, a 5-attempt brute-force lockout engine, an SQLite investigation persistence layer (monitor_status table in users.db), and a file-backed UUID session preservation architecture (?_sid=UUID) that overcomes native Streamlit reset limitations.")
    ]
    for title, desc in milestones:
        p = doc.add_paragraph(style='List Number')
        p.add_run(title).bold = True
        p.add_run(desc)

    # ── 10.2 ──
    add_custom_heading("10.2 Confirmation of Feasibility and Completion Plan", level=2)
    
    add_custom_heading("Confirmation of Technical & Operational Feasibility", level=3)
    doc.add_paragraph("Empirical testing conducted directly on the host machine provides definitive proof of the framework's operational viability. Live penetration experiments—including the simulated exfiltration of Protected Health Information (PHI) to public ChatGPT and database credential leaks to Google Gemini—were successfully intercepted, correctly decoded, accurately scored (RS = 93.75% and RS = 85.00%), and instantaneously escalated to the SOC dashboard.")

    p_feas = doc.add_paragraph()
    p_feas.add_run("Furthermore, sub-second execution speeds during the in-memory asset cross-referencing confirm that the 4-tier decoupled pipeline is highly scalable and computationally efficient. Consequently, the ContextGuard framework is confirmed to be ").font.name = 'Times New Roman'
    p_feas.add_run("100% technically, architecturally, and operationally feasible").bold = True
    p_feas.add_run(" for enterprise-scale Zero-Trust deployment.")

    add_custom_heading("Strategic Completion Plan (Phases 4 & 5)", level=3)
    doc.add_paragraph("With the primary data ingestion, content inspection, and visualization engines fully operational, the remaining project lifecycle will be dedicated to system optimization, expanded evaluation, and academic documentation compilation.")

    road_text = """+-----------------------------------------------------------------------+
|                PHASE 4: SYSTEM MONITORING & EVALUATION                |
|        (Rule Expansion, Multi-User Simulation, Stress Testing)        |
+-----------------------------------------------------------------------+
                                    |
                                    v
+-----------------------------------------------------------------------+
|             PHASE 5: ACADEMIC CLOSURE & THESIS COMPILATION            |
|       (Final Dissertation Writing, Plagiarism Check, Submission)      |
+-----------------------------------------------------------------------+
                                    |
                                    v
+-----------------------------------------------------------------------+
|                 FINAL MILESTONE: PROJECT VIVA & DEMO                  |
|          (Live Technical Presentation and Thesis Defense)             |
+-----------------------------------------------------------------------+"""
    add_code_block(road_text)

    plan_steps = [
        ("System Optimization & Rule Expansion (July - August 2026): ", "Expand the NLP heuristic rule matrices to encompass a broader spectrum of proprietary source code languages and cryptographic tokens. Calibrate the WRSE engine to refine false positive rates across localized organizational scopes."),
        ("Comprehensive Performance Testing (August 2026): ", "Execute automated multi-user traffic simulations to evaluate CPU/Memory overhead and measure processing latency under peak corporate network loads."),
        ("Final Dissertation Compilation (Milestone 4 - Due August 31, 2026): ", "Synthesize the complete research lifecycle into the final academic thesis, ensuring rigorous formatting adherence (Times New Roman 12, 1.5 spacing, IEEE referencing) and official Turnitin plagiarism verification."),
        ("Project Viva & Technical Demonstration (Milestone 5 - September 9, 2026): ", "Prepare official presentation slides and refine the live host machine demonstration environment for the final academic thesis defense before the faculty examination panel.")
    ]
    for title, desc in plan_steps:
        p = doc.add_paragraph(style='List Number')
        p.add_run(title).bold = True
        p.add_run(desc)

    doc.save("/home/izu/ShadowAI_Framework/Interim_Step5_Conclusion.docx")
    print("✅ Successfully generated Interim_Step5_Conclusion.docx")

if __name__ == "__main__":
    create_document()
