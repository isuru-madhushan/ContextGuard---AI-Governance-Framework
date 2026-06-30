import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def create_document():
    doc = docx.Document()

    # ── 1. PAGE MARGINS (1 INCH ALL SIDES) ──
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # ── 2. CONFIGURE STYLES (TIMES NEW ROMAN 12PT, 1.5 SPACING) ──
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = Pt(12)
    normal_style.font.color.rgb = RGBColor(51, 51, 51)
    normal_style.paragraph_format.line_spacing = 1.5
    normal_style.paragraph_format.space_after = Pt(6)

    # Helper function for headings
    def add_custom_heading(text, level):
        p = doc.add_paragraph()
        p.paragraph_format.keep_with_next = True
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.bold = True
        if level == 1:
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(0, 51, 102) # Navy Blue
        elif level == 2:
            run.font.size = Pt(15)
            run.font.color.rgb = RGBColor(0, 102, 153)
        elif level == 3:
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(51, 51, 51)
        return p

    def add_callout(text, caption):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.autofit = False
        tbl.columns[0].width = Inches(6.5)
        cell = tbl.cell(0, 0)
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F4F6F9"/>')
        cell._tc.get_or_add_tcPr().append(shading)
        
        # Borders
        borders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:top w:val="none"/><w:left w:val="single" w:sz="24" w:space="0" w:color="006699"/><w:bottom w:val="none"/><w:right w:val="none"/></w:tcBorders>')
        cell._tc.get_or_add_tcPr().append(borders)
        
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent = Inches(0.1)
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 102, 153)
        
        # Caption below callout
        p_cap = doc.add_paragraph()
        p_cap.paragraph_format.space_before = Pt(4)
        p_cap.paragraph_format.space_after = Pt(12)
        run_cap = p_cap.add_run(caption)
        run_cap.font.name = 'Times New Roman'
        run_cap.italic = True
        run_cap.font.size = Pt(10.5)
        run_cap.font.color.rgb = RGBColor(100, 120, 140)

    # ── BUILD DOCUMENT CONTENT ──
    add_custom_heading("3. Progress Summary", level=1)
    doc.add_paragraph("This section summarizes the tangible engineering milestones achieved up to the interim evaluation stage of the ContextGuard Shadow AI Governance Framework. It details the specific tasks completed, the current operational status of the software prototype, and verifiable evidence of progress, including official source code repository access and live technical video demonstrations.")

    # ── 3.1 ──
    add_custom_heading("3.1 Tasks Completed So Far", level=2)
    doc.add_paragraph("The project has successfully completed the foundation, architectural design, and core prototype development phases (Phases 1 through 3). The specific technical tasks accomplished to date include:")

    tasks = [
        ("Architectural System Design: ", "Finalized the 4-tier decoupled system architecture, separating data ingestion, heuristic/NLP inspection, mathematical risk calculation, and visualization into modular Python components."),
        ("Live Network Interception Engine (live_mitm_logger.py): ", "Developed an active Man-in-the-Middle inspection plane using mitmproxy. Configured targeted HTTP flow filters to intercept outbound TLS-encrypted POST requests directed at prominent AI endpoints (chatgpt.com, gemini.google.com, claude.ai)."),
        ("Advanced Double-Plane URL Decoding Engine: ", "Implemented robust multi-stage parsing algorithms within the logger to decode complex URL entities (f.req=, %5B, %22) from Google Gemini and extract nested JSON arrays from OpenAI/Claude, writing standardized telemetry to wrse_comprehensive_audit.log."),
        ("NLP Normalization & Master Asset Indexing (data_core.py): ", "Engineered forensic_normalize() to eliminate structural payload noise. Established load_master_dataset() to ingest and index 15 Master CSV Data Sheets (data Sheets/) into an in-memory dictionary (idx), categorized by Record IDs, Patient IDs, and high-specificity credential tokens."),
        ("WRSE Mathematical Engine Formulation: ", "Designed and coded the linear weighted sum model calculate_wrse(), establishing pre-calibrated baseline weights for Data Sensitivity (Ws = 0.50), Destination Trust (Wd = 0.25), and User Authority (Wu = 0.25)."),
        ("Zero-Trust Administrative SOC Dashboard (app.py, auth.py, styles.py): ", "Deployed an interactive Streamlit portal featuring WAF input sanitization, 5-attempt brute-force lockout defense, a 4-column top header filter strip, and 1,090+ lines of custom CSS (THREATMON_CSS) for premium glassmorphism styling."),
        ("Stateless Limitations Overcome (Session & SQLite Persistence): ", "Engineered a server-side UUID session persistence engine (/tmp/shadowai_sessions/) synced with URL parameters (?_sid=UUID) to survive manual browser reloads (F5). Built an SQLite persistence layer (components.py) to maintain investigation statuses (monitor_status table in users.db) across hard reloads.")
    ]

    for title, desc in tasks:
        p = doc.add_paragraph(style='List Number')
        p.add_run(title).bold = True
        p.add_run(desc)

    # ── 3.2 ──
    add_custom_heading("3.2 Current Project Status", level=2)
    p = doc.add_paragraph()
    p.add_run("The project is currently at the ").font.name = 'Times New Roman'
    p.add_run("Mid-Implementation and Empirical Validation Stage (Milestone 2 - 100% Complete)").bold = True
    p.add_run(". The primary technical challenges regarding real-time SSL interception, complex URL decoding, memory-efficient CSV indexing, and Streamlit session persistence have been successfully overcome. The core system functions autonomously as a cohesive, operational prototype. The remaining project timeline is dedicated to expanding the rule matrices, optimizing execution latency for large-scale enterprise simulation, conducting extensive boundary testing, and compiling the final research dissertation (Milestones 3 and 4).")

    # ── 3.3 ──
    add_custom_heading("3.3 Evidence of Progress (Artefacts, GitHub Repo & Demo Video)", level=2)
    doc.add_paragraph("To provide verifiable proof of the progress achieved, the complete underlying codebase and an empirical walkthrough video have been compiled for academic review.")

    add_custom_heading("Summary of Developed Software Artefacts:", level=3)
    artefacts = [
        ("live_mitm_logger.py: ", "Active proxy interception script and double-plane URL decoding engine."),
        ("wrse_comprehensive_audit.log: ", "Standardized JSON telemetry log output generated by the interception plane."),
        ("data_core.py: ", "Central NLP parsing, CSV asset indexing, and WRSE mathematical scoring module."),
        ("auth.py: ", "Zero-Trust authentication portal, WAF regex engine, brute-force lockout tracker, and session persistence logic."),
        ("components.py: ", "Frontend layout helpers and SQLite database persistence functions (users.db)."),
        ("app.py & styles.py: ", "Main Streamlit execution script and bespoke CSS glassmorphism injection definitions."),
        ("data Sheets/: ", "The 15 Master CSV corporate asset sheets containing baseline weights and sensitive tokens.")
    ]
    for title, desc in artefacts:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(title).bold = True
        p.add_run(desc)

    add_custom_heading("Official GitHub Repository Link", level=3)
    add_callout("[ INSERT GITHUB REPOSITORY LINK HERE ]", "(Example: https://github.com/yourusername/ShadowAI_ContextGuard_Framework) (*Instructions for Student: Create a public or unlisted GitHub repository, push your active project folder, and paste the exact URL here.*)")

    add_custom_heading("Official Live Demo Video Link", level=3)
    add_callout("[ INSERT YOUTUBE / GOOGLE DRIVE DEMO VIDEO LINK HERE ]", "(Example: https://youtube.com/watch?v=your_video_id or Google Drive shareable link) (*Instructions for Student: Upload your screen recording to YouTube (Unlisted) or Google Drive (Anyone with link can view) and paste the exact link here.*)")

    add_custom_heading("Recommended Step-by-Step Script for Your Demo Video (3–5 Minutes)", level=3)
    doc.add_paragraph("To achieve the highest possible grading during the interim review, follow this exact recording workflow in your demonstration video:")

    steps = [
        ("Step 1 - System Initialization (0:00 - 0:45): ", "Open your terminal and show mitmproxy -s Section1_DataIngestion/live_mitm_logger.py running in the background. Open your browser, navigate to the ContextGuard login screen, and show the premium dark glassmorphism UI."),
        ("Step 2 - Security Defenses Demo (0:45 - 1:30): ", "Type an SQL injection string (' OR 1=1 --) into the login box to trigger and display the custom WAF Security Alert Badge. Log in successfully with valid credentials (admin / adminpass) and demonstrate the main AI Discovery dashboard loading."),
        ("Step 3 - Live Host Machine Exfiltration Experiment (1:30 - 2:45): ", "Open a new browser tab on your host machine and navigate to chatgpt.com or gemini.google.com. Copy a sensitive Medical Record ID (HL7-517169) from your CSV files and paste it into ChatGPT: 'Refactor this HL7 header: HL7-517169'. Submit the prompt. Switch back to your terminal to show [🛡️ INGESTION PLANE CLEANED] confirming live capture. Switch to the ContextGuard dashboard, verify the new threat entry appearing at the top with a 93.75% WRSE Score and a red CRITICAL badge. Click 'Inspect' to show the mathematical calculation breakdown."),
        ("Step 4 - Session & SQLite Persistence Demonstration (2:45 - 3:30): ", "Change the 'Monitor' selectbox for the captured event from Open to In Progress. Hit F5 in your browser to execute a hard page reload. Point out that the URL parameter (?_sid=UUID) successfully preserved your session without logging you out. Select In Progress in the top header filter box and demonstrate that your table status remained perfectly static and filtered, proving SQLite persistence.")
    ]
    for title, desc in steps:
        p = doc.add_paragraph(style='List Number')
        p.add_run(title).bold = True
        p.add_run(desc)

    doc.save("/home/izu/ShadowAI_Framework/Interim_Step4_Progress_Summary.docx")
    print("✅ Successfully generated Interim_Step4_Progress_Summary.docx")

if __name__ == "__main__":
    create_document()
