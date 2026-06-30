import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def create_document():
    doc = docx.Document()

    # ── 1. PAGE MARGINS (1 INCH ALL SIDES) ──
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # ── 2. CONFIGURE STYLES (TIMES NEW ROMAN 12PT, 1.5 SPACING) ──
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = Pt(12)
    normal_style.font.color.rgb = RGBColor(51, 51, 51)
    normal_style.paragraph_format.line_spacing = 1.5
    normal_style.paragraph_format.space_after = Pt(6)

    # Helper function for headings
    def add_custom_heading(text, level):
        p = doc.add_paragraph()
        p.paragraph_format.keep_with_next = True
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.bold = True
        if level == 1:
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(0, 51, 102) # Navy Blue
        elif level == 2:
            run.font.size = Pt(15)
            run.font.color.rgb = RGBColor(0, 102, 153)
        elif level == 3:
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(51, 51, 51)
        return p

    def add_callout(text, caption):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.autofit = False
        tbl.columns[0].width = Inches(6.5)
        cell = tbl.cell(0, 0)
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F4F6F9"/>')
        cell._tc.get_or_add_tcPr().append(shading)
        
        # Borders
        borders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:top w:val="none"/><w:left w:val="single" w:sz="24" w:space="0" w:color="006699"/><w:bottom w:val="none"/><w:right w:val="none"/></w:tcBorders>')
        cell._tc.get_or_add_tcPr().append(borders)
        
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent = Inches(0.1)
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 102, 153)
        
        # Caption below callout
        p_cap = doc.add_paragraph()
        p_cap.paragraph_format.space_before = Pt(4)
        p_cap.paragraph_format.space_after = Pt(12)
        run_cap = p_cap.add_run(caption)
        run_cap.font.name = 'Times New Roman'
        run_cap.italic = True
        run_cap.font.size = Pt(10.5)
        run_cap.font.color.rgb = RGBColor(100, 120, 140)

    def add_code_block(code_text):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.autofit = False
        tbl.columns[0].width = Inches(6.5)
        cell = tbl.cell(0, 0)
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="2D3748"/>') # Dark bg
        cell._tc.get_or_add_tcPr().append(shading)
        
        p = cell.paragraphs[0]
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent = Inches(0.1)
        run = p.add_run(code_text)
        run.font.name = 'Courier New'
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(247, 250, 252) # Light text
        
        doc.add_paragraph() # spacing

    # ── BUILD DOCUMENT CONTENT ──
    add_custom_heading("6. Design and Implementation Progress", level=1)
    doc.add_paragraph("This section outlines the concrete technical progress achieved during the mid-implementation stage of the ContextGuard Shadow AI Governance Framework. It details the underlying system architecture, the active modules and functions developed, and the visual proof of the working prototype.")

    # ── 6.1 ──
    add_custom_heading("6.1 Architecture / Design Diagrams", level=2)
    doc.add_paragraph("The system operates as a transparent inspection layer structured into four decoupled logical tiers. This layered design facilitates asynchronous feature extraction, ensuring that network latency is minimized during real-time risk evaluation.")

    arch_text = """+-----------------------------------------------------------------------+
|                 4. VISUALIZATION & GOVERNANCE LAYER                   |
|       (Streamlit Dashboard, Real-time Alerts, SQLite Persistence)     |
+-----------------------------------------------------------------------+
                                    ^
                                    |  Real-time Risk Telemetry
+-----------------------------------------------------------------------+
|                        3. WRSE SCORING MODULE                         |
|      RS = (Ws * S) + (Wd * D) + (Wu * U)   [Normalized 0-100 Scale]   |
+-----------------------------------------------------------------------+
                                    ^
            +-----------------------+-----------------------+
            | (Behavioral Vectors)                          | (Semantic Vectors)
+---------------------------+               +---------------------------+
|    HEURISTIC ENGINE       |               |        NLP ENGINE         |
|  (IAT & Packet Anomaly)   |               |  (TF-IDF & Asset Match)   |
+---------------------------+               +---------------------------+
            ^                                               ^
            +-----------------------+-----------------------+
                                    |  Decoupled Streams
+-----------------------------------------------------------------------+
|                         1. DATA INGESTION MODULE                      |
|             (live_mitm_logger.py & Synthetic Host Telemetry)          |
+-----------------------------------------------------------------------+"""
    add_code_block(arch_text)

    add_callout("[ INSERT SCREENSHOT 1: Official Architecture Diagram ]", "Caption: Figure 6.1 - The 4-Tier System Architecture of the ContextGuard Shadow AI Governance Framework. (*Instructions for Student: Insert your official architectural diagram image here, illustrating the data flow from Ingestion to the Streamlit Dashboard.*)")

    add_custom_heading("Architectural Tier Breakdown:", level=3)
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Data Ingestion Module: ").bold = True
    p.add_run("Captures live workstation telemetry, API logs, and unencrypted prompt payloads via active Man-in-the-Middle (MITM) interception.")
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Inspection Layer: ").bold = True
    p.add_run("Executes parallel heuristic analysis (evaluating Inter-Arrival Time and packet size variance for bot detection) and deep content inspection (matching prompt payloads against the Master Asset Registry).")
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("WRSE Scoring Module: ").bold = True
    p.add_run("Processes the extracted vectors through the mathematical engine to compute the unified risk coefficient.")
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Visualization & Governance Layer: ").bold = True
    p.add_run("Presents actionable intelligence, time-series filtering, and administrative override options via the Streamlit dashboard.")

    # ── 6.2 ──
    add_custom_heading("6.2 Implementation Progress (Modules Completed, Functions Developed, Prototypes)", level=2)
    doc.add_paragraph("To date, the project has successfully transitioned from conceptual formulation into a fully operational software prototype. The following core modules and functions have been developed and integrated.")

    # ── MODULE 1 ──
    add_custom_heading("Module 1: Host Machine Traffic Interception (live_mitm_logger.py)", level=3)
    doc.add_paragraph("To capture live interactions between enterprise workstations and external AI platforms, we developed a real-time interception module utilizing mitmproxy. This script acts as a transparent inspection layer directly on the host machine, intercepting outbound TLS-encrypted HTTPS traffic before it exits the corporate perimeter.")

    doc.add_paragraph("Developed Functions & Mechanics:").bold = True
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("request(flow) - Targeted Domain Filtering: ").bold = True
    p.add_run("Inspects the HTTPFlow object and matches the destination hostname against a predefined list of prominent generative AI endpoints (chatgpt.com, api.openai.com, gemini.google.com, claude.ai, deepseek.com, copilot.microsoft.com). Only outbound POST requests directed at these services are intercepted for deep packet inspection, preserving network bandwidth and employee privacy for standard web browsing.")

    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Advanced Double-Plane URL Decoding Engine: ").bold = True
    p.add_run("Generative AI platforms employ vastly different payload encoding mechanisms. For instance, Google Gemini frequently encapsulates prompts in complex URL-encoded structures (f.req=), whereas OpenAI and Claude utilize structured JSON bodies. live_mitm_logger.py implements a robust multi-stage parser: Stage 1 (URL Entity Decoding) detects URL-encoded strings and applies urllib.parse.unquote to strip garbage boundary structures. Stage 2 (JSON Structure Parsing) safely parses JSON bodies to extract raw text from deeply nested keypaths.")

    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Structured Telemetry Archiving: ").bold = True
    p.add_run("The extracted prompt payload is combined with critical connection metadata—including Client IP, Source Port, Destination Domain, Destination IP, HTTP Method, User-Agent, and exact timestamps. This unified telemetry object is dynamically appended to wrse_comprehensive_audit.log in structured JSON format.")

    add_callout("[ INSERT SCREENSHOT 2: Code Segment from live_mitm_logger.py (Lines 14 to 48) ]", "Caption: Figure 6.2 - Code implementation of the request interception and Advanced Double-Plane URL Decoding Engine in live_mitm_logger.py (Lines 14–48). (*Instructions for Student: Open Section1_DataIngestion/live_mitm_logger.py in VS Code, scroll to lines 14–48 showing the request function and decoding logic, and take a clear screenshot.*)")

    add_callout("[ INSERT SCREENSHOT 3: JSON Log Structure from wrse_comprehensive_audit.log ]", "Caption: Figure 6.3 - The structured JSON log format within wrse_comprehensive_audit.log containing connection metadata, source/dest nodes, and the extracted raw prompt. (*Instructions for Student: Open Section1_DataIngestion/wrse_comprehensive_audit.log in VS Code, highlight a clean JSON log entry, and take a screenshot.*)")

    # ── MODULE 2 ──
    add_custom_heading("Module 2: Payload Extraction, Normalization & Master Asset Indexing (data_core.py)", level=3)
    doc.add_paragraph("Once the raw logs are written to disk, the central analytical engine (data_core.py) executes a highly sophisticated pre-processing and indexing routine to prepare the unstructured prompts for mathematical risk evaluation.")

    doc.add_paragraph("Developed Functions & Mechanics:").bold = True

    p = doc.add_paragraph(style='List Bullet')
    p.add_run("forensic_normalize(text): ").bold = True
    p.add_run("Unstructured prompts frequently contain escape characters, quotes, and nested array brackets that disrupt standard NLP string matching. This function applies a multi-pass regex filter to extract pure text strings, strip out structural noise, and convert all characters to lowercase while preserving critical punctuation.")

    p = doc.add_paragraph(style='List Bullet')
    p.add_run("load_master_dataset() - Master 15-Sheet CSV Indexing: ").bold = True
    p.add_run("To enable context-aware inspection, the framework ingests 15 Master CSV Data Sheets located in data Sheets/. These sheets represent organizational assets across various tiers (Medical Records/PHI, Infrastructure Core Assets, Intellectual Property, Financial Data, HR Records). To achieve sub-second matching speeds without disk I/O bottlenecks, this function builds a highly optimized in-memory dictionary index (idx), categorizing assets by record_ids (IAM-670469, HL7-517169), patient_ids (TM-XXXX-XXXX), and high-specificity tokens (SAML Token Hashes, API Key Hashes, DB Connection Strings).")

    p = doc.add_paragraph(style='List Bullet')
    p.add_run("find_master_asset_match(prompt_text, asset_index): ").bold = True
    p.add_run("Scans the text for explicit Record/Patient IDs using regex and cross-references them against the token index. It utilizes a seen_records set and matched_token_values tracking to ensure that if a prompt contains both an Asset ID and its corresponding Database String, it is correctly counted as a single exposed asset rather than creating artificial mock collisions.")

    add_callout("[ INSERT SCREENSHOT 4: Code Segment from data_core.py (Lines 232 to 272) ]", "Caption: Figure 6.4 - Implementation of the find_master_asset_match function in data_core.py (Lines 232–272), showcasing token matching and collision defense logic. (*Instructions for Student: Open Section3_Dashboard/data_core.py in VS Code, scroll to lines 232–272 showing the find_master_asset_match function, and take a screenshot.*)")

    add_callout("[ INSERT SCREENSHOT 5: Master CSV Asset Registry Indexing in File Explorer ]", "Caption: Figure 6.5 - The 15 Master CSV data sheets located in the 'data Sheets' directory, containing baseline weights and high-specificity tokens. (*Instructions for Student: Take a screenshot of the VS Code file explorer showing the list of T1_xxx.csv, T2_xxx.csv files inside the 'data Sheets' folder.*)")

    # ── MODULE 3 ──
    add_custom_heading("Module 3: Real-Time WRSE Risk Calculation & Severity Mapping (data_core.py)", level=3)
    doc.add_paragraph("The ultimate technical objective of ContextGuard is translating the extracted metadata and asset matches into a standardized, actionable Risk Score (RS). This is handled by calculate_wrse().")

    doc.add_paragraph("Developed Functions & Mechanics:").bold = True

    p = doc.add_paragraph(style='List Bullet')
    p.add_run("calculate_wrse(prompt_text, dest_trust_w, user_auth_w, asset_matches): ").bold = True
    p.add_run("Accepts the extracted prompt, destination penalty, user privilege weight, and matched asset records. It applies pre-calibrated baseline constants: Ws = 0.50 (Sensitivity), Wd = 0.25 (Destination Trust), and Wu = 0.25 (User Authority).")

    p_eq = doc.add_paragraph()
    p_eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_eq = p_eq.add_run("RS = (Ws × S) + (Wd × D) + (Wu × U)")
    run_eq.bold = True
    run_eq.font.size = Pt(14)
    run_eq.font.color.rgb = RGBColor(0, 51, 102)
    doc.add_paragraph() # space

    p = doc.add_paragraph(style='List Bullet')
    p.add_run("get_severity(score): ").bold = True
    p.add_run("Evaluates the normalized 0–100 score against strict administrative thresholds: Score > 80 -> CRITICAL (🔴), Score >= 55 -> MEDIUM (🟡), and Score < 55 -> LOW (🟢).")

    doc.add_paragraph("Detailed Numerical Walkthrough of an Active Threat:").bold = True
    doc.add_paragraph("Consider a real-world scenario processed by the prototype: A software engineer (idx_e = 0) copies an internal HL7 healthcare protocol header (HL7-517169) and pastes it into chatgpt.com to request a code refactor.")

    p = doc.add_paragraph(style='List Number')
    p.add_run("Step 1 (Sensitivity Score S Calculation): ").bold = True
    p.add_run("The NLP engine matches HL7-517169 against the Master Asset Registry. The CSV definition establishes a baseline Asset Weight Wi = 0.95 (Tier 1 - Medical Records/PHI). S = min(∑ Wi, 1.0) = min(0.95, 1.0) = 0.95.")

    p = doc.add_paragraph(style='List Number')
    p.add_run("Step 2 (Destination Trust D Assignment): ").bold = True
    p.add_run("process_events() inspects dest_domain. Because chatgpt.com matches the unmanaged public LLM list (gemini, chatgpt, claude, openai), the destination penalty maximizes: D = 0.95.")

    p = doc.add_paragraph(style='List Number')
    p.add_run("Step 3 (User Authority U Assignment): ").bold = True
    p.add_run("Based on the user flow index (idx_e % 2 == 0), the user privilege weight is assigned: U = 0.90.")

    p = doc.add_paragraph(style='List Number')
    p.add_run("Step 4 (Applying the Linear Weighted Sum Model): ").bold = True
    p.add_run("The formula executes: RS = (Ws × S) + (Wd × D) + (Wu × U) = (0.50 × 0.95) + (0.25 × 0.95) + (0.25 × 0.90) = 0.475 + 0.2375 + 0.225 = 0.9375.")

    p = doc.add_paragraph(style='List Number')
    p.add_run("Step 5 (Normalization & Severity Mapping): ").bold = True
    p.add_run("The raw coefficient is scaled to a 100-point index: round(0.9375 * 100, 2) = 93.75%. The get_severity(93.75) function evaluates the threshold (Score > 80). It instantly attaches a CRITICAL label and a red visual alert icon (🔴), pushing the event to the top of the SOC dashboard feed.")

    add_callout("[ INSERT SCREENSHOT 6: Code Segment from data_core.py (Lines 88 to 117) ]", "Caption: Figure 6.6 - Implementation of the calculate_wrse function in data_core.py (Lines 88–117), showing the linear weighted sum model execution. (*Instructions for Student: Open Section3_Dashboard/data_core.py in VS Code, scroll to lines 88–117 showing calculate_wrse, and take a screenshot.*)")

    add_callout("[ INSERT SCREENSHOT 7: Real-time WRSE Score Breakdown in Dashboard UI ]", "Caption: Figure 6.7 - Real-time dashboard display of a captured Shadow AI event showcasing the 93.75% WRSE score, triggered asset keys, and CRITICAL severity badge. (*Instructions for Student: Take a screenshot of the dashboard table row or the 'Inspect' detail view showing a high WRSE score and its calculation breakdown.*)")

    # ── MODULE 4 ──
    add_custom_heading("Module 4: Secure Authentication & Session Persistence (auth.py)", level=3)
    doc.add_paragraph("The authentication system was designed to mirror enterprise-grade Zero-Trust access portals.")

    doc.add_paragraph("Developed Functions & Mechanics:").bold = True
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("sanitize_input(val) - Web Application Firewall (WAF) Sanitization: ").bold = True
    p.add_run("Incorporated active regex checks to intercept and block SQL injection (' OR 1=1 --) and XSS (<script>) payloads at the login prompt.")

    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Brute-Force Lockout Defense: ").bold = True
    p.add_run("Implemented a stateful lockout mechanism that freezes the login interface for 30 seconds upon registering 5 consecutive invalid login attempts.")

    p = doc.add_paragraph(style='List Bullet')
    p.add_run("create_session, validate_session, check_auth - File-Backed Session Persistence: ").bold = True
    p.add_run("Solved Streamlit's native session-reset behavior on browser refreshes by engineering a custom file-backed token mechanism (/tmp/shadowai_sessions/). Upon successful authentication, a unique UUID session token is generated, stored on the server, and injected into the client's URL query parameters (?_sid=UUID). When a user presses F5 or refreshes the tab, check_auth() intercepts the URL parameter, validates the active session file, and seamlessly restores the user state without forcing a re-login.")

    add_callout("[ INSERT SCREENSHOT 8: Code Segment from auth.py (Lines 466 to 485) ]", "Caption: Figure 6.8 - Code implementation of persistent session token creation and URL query parameter storage upon successful authentication in auth.py (Lines 466–485). (*Instructions for Student: Open Section3_Dashboard/auth.py in VS Code, scroll to lines 466–485 showing the 'if success:' block and create_session call, and take a screenshot.*)")

    add_callout("[ INSERT SCREENSHOT 9: Premium Glassmorphism Login Page UI ]", "Caption: Figure 6.9 - The ContextGuard premium glassmorphism login interface featuring bespoke styling, embedded branding, and active WAF sanitization. (*Instructions for Student: Take a screenshot of the login page showing the dark background, logo, and login card.*)")

    add_callout("[ INSERT SCREENSHOT 10: Brute-Force Lockout Alert UI ]", "Caption: Figure 6.10 - Active brute-force protection triggering a 30-second administrative lockout after 5 failed login attempts. (*Instructions for Student: Enter wrong passwords 5 times and take a screenshot of the red lockout error message.*)")

    # ── MODULE 5 ──
    add_custom_heading("Module 5: Stateful UI & Event Monitoring (app.py, components.py, styles.py)", level=3)
    doc.add_paragraph("The main administrative interface has been fully constructed using a clean, non-scrolling, high-density layout.")

    doc.add_paragraph("Developed Functions & Mechanics:").bold = True
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("app.py - Top Header Filtering Controls: ").bold = True
    p.add_run("Implemented a sophisticated 4-column header bar allowing administrators to instantly filter threat feeds by Time Type, Time Interval, and Monitor Status (All Statuses, Open, In Progress, Close).")

    p = doc.add_paragraph(style='List Bullet')
    p.add_run("components.py - SQLite-Backed State Persistence: ").bold = True
    p.add_run("Developed backend helper functions (_init_monitor_table, _get_monitor_status, _save_monitor_status) that instantly write status updates from the table's interactive selectbox to a dedicated monitor_status table in users.db. When an administrator updates an event's status from Open to In Progress, the change is immediately committed to the database. Consequently, the status remains perfectly static and preserved across automatic background refreshes and manual browser reloads.")

    p = doc.add_paragraph(style='List Bullet')
    p.add_run("styles.py - Custom UI Styling & Glassmorphism: ").bold = True
    p.add_run("Injected over 1,090 lines of custom CSS (THREATMON_CSS) to completely restyle native BaseWeb containers, implement micro-animations, add glowing risk borders, and suppress native Streamlit branding (headers, footers, running spinners).")

    add_callout("[ INSERT SCREENSHOT 11: Code Segment from components.py (Lines 204 to 215) ]", "Caption: Figure 6.11 - Code implementation of the interactive Monitor status selectbox inside the event table rendering logic in components.py (Lines 204–215). (*Instructions for Student: Open Section3_Dashboard/components.py in VS Code, scroll to lines 204–215 showing the selectbox configuration, and take a screenshot.*)")

    add_callout("[ INSERT SCREENSHOT 12: Main AI Discovery Dashboard UI with Top Header Filters ]", "Caption: Figure 6.12 - The primary AI Discovery dashboard view showcasing the top metric strip, 4-column header filter controls, and the active threat table. (*Instructions for Student: Take a full-screen screenshot of the main dashboard page showing the top strip, filter boxes, and data table.*)")

    add_callout("[ INSERT SCREENSHOT 13: Sidebar Navigation UI with Relocated Sign Out Button ]", "Caption: Figure 6.13 - The restructured sidebar navigation panel displaying the active monitoring badge, relocated Sign Out button, and logged-in user credentials. (*Instructions for Student: Take a clear screenshot of the left sidebar showing the logo, Sign Out button, and user role badge.*)")

    add_callout("[ INSERT SCREENSHOT 14: Filtered 'In Progress' Threat Feed UI ]", "Caption: Figure 6.14 - The dashboard dynamically filtered to display only events assigned the 'In Progress' monitoring status via the top header control. (*Instructions for Student: Select 'In Progress' in the top header filter box and take a screenshot showing the filtered results.*)")

    doc.save("/home/izu/ShadowAI_Framework/Interim_Step1_Design_Implementation.docx")
    print("✅ Successfully generated Interim_Step1_Design_Implementation.docx")

if __name__ == "__main__":
    create_document()
