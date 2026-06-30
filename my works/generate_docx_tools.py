import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def create_document():
    doc = docx.Document()

    # ── 1. PAGE MARGINS (1 INCH ALL SIDES) ──
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # ── 2. CONFIGURE STYLES (TIMES NEW ROMAN 12PT, 1.5 SPACING) ──
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = Pt(12)
    normal_style.font.color.rgb = RGBColor(51, 51, 51)
    normal_style.paragraph_format.line_spacing = 1.5
    normal_style.paragraph_format.space_after = Pt(6)

    # Helper function for headings
    def add_custom_heading(text, level):
        p = doc.add_paragraph()
        p.paragraph_format.keep_with_next = True
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.bold = True
        if level == 1:
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(0, 51, 102) # Navy Blue
        elif level == 2:
            run.font.size = Pt(15)
            run.font.color.rgb = RGBColor(0, 102, 153)
        elif level == 3:
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(51, 51, 51)
        return p

    add_custom_heading("5.3 Tools, Techniques, Technologies Selected and Justification", level=1)
    doc.add_paragraph("The operational transition from a theoretical framework into an enterprise-grade Zero-Trust prototype required an expanded technical stack. In addition to the foundational data science libraries, several specialized network interception, database persistence, and cryptographic modules were integrated into the production environment.")

    add_custom_heading("Expanded Technical Stack & Library Justification", level=2)

    table_data = [
        ("Core Engine", "Python 3.10+", "Primary programming language chosen for its extensive asynchronous networking capabilities, robust standard libraries, and seamless integration with ML/NLP frameworks."),
        ("Network Interception", "mitmproxy (mitmproxy.http)", "Core proxy kernel utilized in live_mitm_logger.py to establish a Man-in-the-Middle inspection plane on the host machine. It transparently intercepts and exposes outbound TLS-encrypted HTTPS POST flows directed at unmanaged AI endpoints (chatgpt.com, gemini.google.com)."),
        ("Payload Decoding", "urllib.parse & json", "Foundational modules powering the Advanced Double-Plane URL Decoding Engine. urllib.parse.unquote strips percent-encoded entities (%5B, %22, f.req=) from Google Gemini prompts, while json.loads extracts deeply nested prompt arrays from OpenAI payloads."),
        ("Data Processing & Indexing", "Pandas & NumPy", "Utilized in data_core.py for high-performance in-memory dataset manipulation, statistical aggregation of network metadata, and rapid loading of the 15 Master CSV asset sheets."),
        ("NLP & Deep Inspection", "NLTK & re (Regex Engine)", "Serves as the central intelligence engine for the content inspection layer. The re module executes multi-pass regex normalization (forensic_normalize), while NLTK handles tokenization and TF-IDF weighting for the Sensitivity Score (S)."),
        ("Database Persistence", "SQLite3 (sqlite3)", "Lightweight, serverless relational database engine utilized in components.py and auth.py (users.db). It provides atomic state persistence for the monitor_status table, ensuring investigation states (Open, In Progress, Close) remain locked across hard browser reloads."),
        ("Zero-Trust Security", "uuid & Server Storage (/tmp)", "Powers the bespoke session persistence architecture in auth.py. uuid.uuid4() generates cryptographically secure session identification tokens (_sid) stored in /tmp/shadowai_sessions/, preserving administrative sessions across page refreshes (F5)."),
        ("Dashboarding & UI", "Streamlit", "Rapid frontend deployment framework providing the administrative Security Operations Center (SOC) interface, real-time threat feed rendering, and visual telemetry charts."),
        ("Custom Aesthetics", "HTML5 & Custom CSS Injection", "Over 1,090 lines of bespoke CSS (THREATMON_CSS in styles.py) injected via st.markdown(unsafe_allow_html=True). It overrides native BaseWeb containers to implement premium glassmorphism panels, dark backdrops, glowing risk borders, and micro-animations."),
        ("IDE & Version Control", "VS Code & Git", "Integrated development environment and versioning infrastructure utilized for modular scripting, debugging, and maintaining an immutable audit trail of source code changes.")
    ]

    table = doc.add_table(rows=1, cols=3)
    table.autofit = False
    table.columns[0].width = Inches(1.3)
    table.columns[1].width = Inches(1.7)
    table.columns[2].width = Inches(3.5)

    # Header Row
    hdr_cells = table.rows[0].cells
    hdr_titles = ["Category", "Tool / Library", "Technical Purpose & Justification"]
    for i, title in enumerate(hdr_titles):
        cell = hdr_cells[i]
        cell.text = title
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="003366"/>') # Navy bg
        cell._tc.get_or_add_tcPr().append(shading)
        p = cell.paragraphs[0]
        p.runs[0].font.name = 'Times New Roman'
        p.runs[0].font.bold = True
        p.runs[0].font.color.rgb = RGBColor(255, 255, 255) # White text
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)

    # Data Rows
    for row_data in table_data:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            cell = row_cells[i]
            cell.text = val
            p = cell.paragraphs[0]
            p.runs[0].font.name = 'Times New Roman'
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            if i < 2:
                p.runs[0].font.bold = True
            
            # Shading alternating or light borders
            borders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/><w:left w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/><w:bottom w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/><w:right w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/></w:tcBorders>')
            cell._tc.get_or_add_tcPr().append(borders)

    doc.add_paragraph() # space

    doc.save("/home/izu/ShadowAI_Framework/Interim_Updated_Tools_Table.docx")
    print("✅ Successfully generated Interim_Updated_Tools_Table.docx")

if __name__ == "__main__":
    create_document()
