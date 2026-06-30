import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def create_document():
    doc = docx.Document()

    # ── 1. PAGE MARGINS (1 INCH ALL SIDES) ──
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # ── 2. CONFIGURE STYLES (TIMES NEW ROMAN 12PT, 1.5 SPACING) ──
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = Pt(12)
    normal_style.font.color.rgb = RGBColor(51, 51, 51)
    normal_style.paragraph_format.line_spacing = 1.5
    normal_style.paragraph_format.space_after = Pt(6)

    # Helper function for headings
    def add_custom_heading(text, level):
        p = doc.add_paragraph()
        p.paragraph_format.keep_with_next = True
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.bold = True
        if level == 1:
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(0, 51, 102) # Navy Blue
        elif level == 2:
            run.font.size = Pt(15)
            run.font.color.rgb = RGBColor(0, 102, 153)
        elif level == 3:
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(51, 51, 51)
        return p

    def add_callout(text, caption):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.autofit = False
        tbl.columns[0].width = Inches(6.5)
        cell = tbl.cell(0, 0)
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F4F6F9"/>')
        cell._tc.get_or_add_tcPr().append(shading)
        
        # Borders
        borders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:top w:val="none"/><w:left w:val="single" w:sz="24" w:space="0" w:color="006699"/><w:bottom w:val="none"/><w:right w:val="none"/></w:tcBorders>')
        cell._tc.get_or_add_tcPr().append(borders)
        
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent = Inches(0.1)
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 102, 153)
        
        # Caption below callout
        p_cap = doc.add_paragraph()
        p_cap.paragraph_format.space_before = Pt(4)
        p_cap.paragraph_format.space_after = Pt(12)
        run_cap = p_cap.add_run(caption)
        run_cap.font.name = 'Times New Roman'
        run_cap.italic = True
        run_cap.font.size = Pt(10.5)
        run_cap.font.color.rgb = RGBColor(100, 120, 140)

    # ── BUILD DOCUMENT CONTENT ──
    add_custom_heading("8. Challenges and Risk Management", level=1)
    doc.add_paragraph("This section provides a transparent evaluation of the technical hurdles, time constraints, resource limitations, and data parsing complexities encountered during the development of the ContextGuard Shadow AI Governance Framework. It details the precise software engineering actions taken to overcome these obstacles and outlines a comprehensive risk mitigation plan for the remaining phases of the project.")

    # ── 8.1 ──
    add_custom_heading("8.1 Issues Faced & Actions Taken to Resolve Challenges", level=2)
    doc.add_paragraph("Developing an enterprise-grade Zero-Trust inspection portal on top of a rapid prototyping framework like Streamlit introduced significant architectural challenges. The following subsections detail the primary obstacles encountered and the technical solutions implemented.")

    # ── CHALLENGE 1 ──
    add_custom_heading("Challenge 1: Streamlit Session State Volatility on Browser Refreshes", level=3)
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Issue Faced (Technical & Architectural): ").bold = True
    p.add_run("Streamlit operates on a stateless execution model where the entire Python script re-runs from top to bottom upon every user interaction or page reload (F5). During initial prototype testing, this behavior caused active administrative sessions to completely reset whenever the browser was refreshed, instantly logging users out and returning them to the login screen.")

    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Action Taken to Resolve: ").bold = True
    p.add_run("We engineered a custom server-side session persistence architecture within auth.py. Upon successful authentication, the system generates a cryptographically secure UUID session token, stores the session metadata in a temporary server-side directory (/tmp/shadowai_sessions/), and dynamically injects the token into the client's browser URL query parameters (?_sid=UUID). When an administrator refreshes the tab, check_auth() intercepts the URL parameter, executes validate_session(), verifies the local session file, and seamlessly restores st.session_state['authenticated'] = True without disrupting the user experience.")

    add_callout("[ INSERT SCREENSHOT 1: Code Segment from auth.py (Lines 115 to 135) ]", "Caption: Figure 8.1 - Code implementation of the validate_session function in auth.py (Lines 115–135), managing file-backed session verification via URL query parameters. (*Instructions for Student: Open Section3_Dashboard/auth.py in VS Code, scroll to lines 115–135 showing validate_session, and take a screenshot.*)")

    # ── CHALLENGE 2 ──
    add_custom_heading("Challenge 2: UI State Loss in BaseWeb Table Component & Monitor Status Persistence", level=3)
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Issue Faced (Technical & UI Limitations): ").bold = True
    p.add_run("To enable SOC administrators to track threat investigations, we integrated an interactive selectbox (Open, In Progress, Close) directly into the event monitoring table. However, because Streamlit dataframes are inherently stateless, any time an administrator updated an event's status, the table would reset to its default state upon the next background telemetry refresh or manual browser reload.")

    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Action Taken to Resolve: ").bold = True
    p.add_run("We decoupled the table UI state from the frontend dataframe by engineering an SQLite-backed state persistence layer in components.py. We initialized a dedicated monitor_status table inside users.db. Whenever an administrator toggles a selectbox, the helper function _save_monitor_status() instantly commits the new status (EventID, Status) to the database. During table rendering, _get_monitor_status() queries the database and dynamically updates df_all, ensuring that investigation statuses remain perfectly static and preserved across all reloads.")

    add_callout("[ INSERT SCREENSHOT 2: Code Segment from components.py (Lines 18 to 44) ]", "Caption: Figure 8.2 - Code implementation of the SQLite state persistence functions (_save_monitor_status and _get_monitor_status) in components.py (Lines 18–44). (*Instructions for Student: Open Section3_Dashboard/components.py in VS Code, scroll to lines 18–44 showing the SQLite helper functions, and take a screenshot.*)")

    # ── CHALLENGE 3 ──
    add_custom_heading("Challenge 3: Advanced Custom UI Styling & Glassmorphism over Streamlit BaseWeb", level=3)
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Issue Faced (UI Aesthetics & Framework Limitations): ").bold = True
    p.add_run("Native Streamlit applications exhibit a basic, standardized visual appearance that fails to convey the premium, dark-themed, glassmorphism aesthetics expected of a modern enterprise Security Operations Center (SOC) dashboard. Streamlit does not natively support deep DOM manipulation or customized CSS styling for individual BaseWeb containers.")

    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Action Taken to Resolve: ").bold = True
    p.add_run("We bypassed Streamlit's native styling engine by creating a comprehensive custom styling module (styles.py). We injected over 1,090 lines of highly specific CSS (THREATMON_CSS) utilizing st.markdown(..., unsafe_allow_html=True). This custom stylesheet restyles BaseWeb containers with semi-transparent dark backgrounds (background: rgba(23, 27, 38, 0.75)), applies backdrop blur filters (backdrop-filter: blur(16px)), adds glowing neon borders for critical risks, implements subtle micro-animations on hover, and suppresses native Streamlit branding elements (headers, footers, and execution spinners).")

    add_callout("[ INSERT SCREENSHOT 3: Code Segment from styles.py (Lines 15 to 65) ]", "Caption: Figure 8.3 - Code snippet of the custom THREATMON_CSS definitions in styles.py (Lines 15–65), injecting custom glassmorphism properties and overriding BaseWeb containers. (*Instructions for Student: Open Section3_Dashboard/styles.py in VS Code, scroll to lines 15–65 showing the CSS definitions, and take a screenshot.*)")

    # ── CHALLENGE 4 ──
    add_custom_heading("Challenge 4: Data Limitations & NLP Parsing (Google Gemini f.req URL Encoding vs OpenAI JSON)", level=3)
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Issue Faced (Data Limitations & Interception Complexity): ").bold = True
    p.add_run("Extracting pure prompt text from intercepted HTTPS POST requests proved highly complex because external generative AI platforms utilize fundamentally divergent payload architectures. While OpenAI (api.openai.com) and Claude (claude.ai) transmit clean, structured JSON arrays, Google Gemini (gemini.google.com) encapsulates prompts within complex, deeply nested URL-encoded string structures (f.req=%5B%22...%22%5D). Standard string matching failed completely on Gemini payloads due to severe structural noise.")

    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Action Taken to Resolve: ").bold = True
    p.add_run("We developed the Advanced Double-Plane URL Decoding Engine within live_mitm_logger.py. The engine actively inspects the raw payload structure. If it detects f.req= or URL entities (%5B, %22), it routes the payload through a multi-pass urllib.parse.unquote routine, strips out the surrounding boundary garbage, and passes a perfectly clean text string to data_core.py for NLP evaluation.")

    # ── CHALLENGE 5 ──
    add_custom_heading("Challenge 5: Time & Computational Resources (Master Asset Registry Collision & Memory Overhead)", level=3)
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Issue Faced (Time, Resources & Disk I/O Bottlenecks): ").bold = True
    p.add_run("To establish Zero-Trust data governance, ContextGuard must scan every intercepted prompt against 15 Master CSV Asset Sheets (data Sheets/) representing Medical Records, Infrastructure Credentials, Financial Data, and IP. During initial testing, repeatedly opening and parsing 15 CSV files from disk for every incoming network packet created severe disk I/O bottlenecks, raising response times to over 4 seconds and causing false-positive mock collisions.")

    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Action Taken to Resolve: ").bold = True
    p.add_run("We refactored data_core.py to eliminate real-time disk I/O. Using load_master_dataset(), the framework ingests and indexes all 15 CSV sheets into a highly optimized in-memory dictionary (idx) during system startup. Furthermore, we integrated a seen_records set and matched_token_values tracking inside find_master_asset_match(). This algorithmic optimization ensures sub-second evaluation speeds while preventing duplicate collision counting when a single prompt contains multiple related asset parameters.")

    # ── 8.2 RISK MITIGATION MATRIX ──
    add_custom_heading("8.2 Risk Mitigation Plan", level=2)
    doc.add_paragraph("To ensure the long-term viability and successful final delivery of the ContextGuard framework, we have established a proactive risk management matrix. This plan anticipates potential operational, technical, and external risks during the final implementation phase.")

    # Create Table
    table_data = [
        ("RM-01", "External AI API Schema Changes", "High", "High", "External platforms (OpenAI, Gemini, Claude) frequently update their internal POST payload structures. We are implementing an abstract parsing interface in live_mitm_logger.py with fallback regex pattern matching to ensure continuous prompt extraction even if JSON keypaths change."),
        ("RM-02", "mitmproxy Certificate Pinning Bypasses", "High", "Medium", "Enterprise desktop client apps (e.g., native ChatGPT desktop app) may utilize strict SSL certificate pinning, bypassing standard mitmproxy inspection. Mitigation involves deploying custom Enterprise Root CAs via Microsoft Intune / Group Policy Objects (GPO) across workstation trust stores."),
        ("RM-03", "Prompt Injection & Obfuscation Attacks", "Medium", "Medium", "Malicious insiders may attempt to bypass the NLP engine by obfuscating sensitive tokens (e.g., base64 encoding or inserting special characters into Medical IDs). We are expanding forensic_normalize() to execute automated base64 decoding and advanced string normalization prior to asset matching."),
        ("RM-04", "Scalability & Memory Overhead in Large SOCs", "Medium", "Low", "Scaling the in-memory master asset dictionary to accommodate millions of corporate records could increase server RAM usage. Mitigation involves migrating the in-memory dictionary to a dedicated Redis caching layer or an in-memory SQLite database instance for enterprise production deployment."),
        ("RM-05", "Strict Project Timeline & Testing Constraints", "Low", "Low", "Unforeseen debugging requirements could impact the final submission schedule. We have established a strict modular milestone schedule, finalizing core ingestion and UI engines early to reserve the final 4 weeks exclusively for empirical testing and documentation compilation.")
    ]

    table = doc.add_table(rows=1, cols=5)
    table.autofit = False
    table.columns[0].width = Inches(0.8)
    table.columns[1].width = Inches(1.5)
    table.columns[2].width = Inches(0.8)
    table.columns[3].width = Inches(0.8)
    table.columns[4].width = Inches(2.6)

    # Header Row
    hdr_cells = table.rows[0].cells
    hdr_titles = ["Risk ID", "Identified Risk", "Potential Impact", "Probability", "Proposed Mitigation Strategy"]
    for i, title in enumerate(hdr_titles):
        cell = hdr_cells[i]
        cell.text = title
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="003366"/>') # Navy bg
        cell._tc.get_or_add_tcPr().append(shading)
        p = cell.paragraphs[0]
        p.runs[0].font.name = 'Times New Roman'
        p.runs[0].font.bold = True
        p.runs[0].font.color.rgb = RGBColor(255, 255, 255) # White text
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)

    # Data Rows
    for row_data in table_data:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            cell = row_cells[i]
            cell.text = val
            p = cell.paragraphs[0]
            p.runs[0].font.name = 'Times New Roman'
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            if i < 4:
                p.runs[0].font.bold = True
            
            # Shading alternating or light borders
            borders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/><w:left w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/><w:bottom w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/><w:right w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/></w:tcBorders>')
            cell._tc.get_or_add_tcPr().append(borders)

    doc.add_paragraph() # space

    doc.save("/home/izu/ShadowAI_Framework/Interim_Step3_Challenges_RiskManagement.docx")
    print("✅ Successfully generated Interim_Step3_Challenges_RiskManagement.docx")

if __name__ == "__main__":
    create_document()
